#!/usr/bin/env python3
"""
Build a formatted Word document from the humanized markdown manuscript.
Matches the formatting of the original paper2_ai_divide.docx.

Usage:
    python build_docx.py

Source:  manuscripts/paper2_ai_divide_humanized.md
Output:  manuscripts/paper2_ai_divide_humanized.docx
"""

import os
import re
import sys
import io
import copy
from collections import OrderedDict

from docx import Document
from docx.shared import Pt, Inches, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Paths ────────────────────────────────────────────────────────────
BASE = "/Users/hosung/AI_Polarization_Pew/manuscripts/paper2"
MD_PATH = os.path.join(BASE, "manuscripts", "paper2_ai_divide_humanized.md")
ORIG_DOCX = os.path.join(BASE, "manuscripts", "paper2_ai_divide.docx")
OUT_PATH = os.path.join(BASE, "manuscripts", "paper2_ai_divide_humanized.docx")
FIG_DIR = os.path.join(BASE, "output", "figures")

# Figure file mapping (same as original generate_docx.py)
FIGURES = {
    1: ("fig_lca_profiles_form_a.png",
        "Figure 1. Class-conditional response probabilities for the four-class LCA solution (Form A). "
        "Bars represent the probability of each response category within each latent class. "
        "Top indicators (attitude, awareness) are shared across forms; bottom indicators (domains a\u2013d) are Form A-specific."),
    2: ("fig_sem01_ame_by_class.png",
        "Figure 2. Average marginal effects of key predictors on latent class membership probability (full model). "
        "Points represent AMEs with 95% confidence intervals. Reference categories: College graduate+, "
        "Middle income, White NH, Ages 30\u201349, Male, Democrat/Lean Dem."),
    3: ("fig_sem02_mediation_diagram.png",
        "Figure 3. Mediation path diagram: SES to AI awareness to latent class membership. "
        "Standardized coefficients shown for the AI-Advantaged class outcome. "
        "Dashed lines represent indirect paths; solid lines represent direct paths."),
    4: ("fig_sem03_predicted_probs.png",
        "Figure 4. Predicted probability of latent class membership by education level and race/ethnicity. "
        "Predictions from the full multinomial logistic model, holding other variables at reference/median values."),
    5: ("fig_cv01_profile_comparison.png",
        "Figure 5. Cross-wave profile comparison: class-conditional probabilities on shared indicators "
        "(attitude and awareness) across W119 (December 2022), W132 (August 2023), and W152 (August 2024)."),
    6: ("fig_cv03_education_gradient.png",
        "Figure 6. Education gradient in AI-Optimistic/Engaged class membership across survey waves. "
        "Proportion of each education group assigned to the most engaged class, with 95% confidence intervals."),
}

# ── Formatting constants (EMU values from original) ──────────────────
PAGE_WIDTH = Emu(7772400)    # Letter size width
PAGE_HEIGHT = Emu(10058400)  # Letter size height
MARGIN = Emu(914400)         # 1 inch
BODY_INDENT = Emu(457200)    # 0.5 inch first-line indent
HEADING_SPACE = Emu(152400)  # Space before headings (12pt)
AUTHOR_SPACE = Emu(304800)   # Extra space for author info (24pt)
LINE_SPACING = 2.0           # Double-spaced

# ── Original docx resources ─────────────────────────────────────────
_orig_doc = None
_orig_images = None  # {fig_num: blob}
_orig_tables = None


def load_original_docx():
    """Load the original docx and extract images and tables.
    Images are mapped by their position in the document (Figure 1 = first image, etc.)
    """
    global _orig_doc, _orig_images, _orig_tables
    try:
        _orig_doc = Document(ORIG_DOCX)

        # Extract images in document order by scanning paragraphs for inline shapes
        _orig_images = {}
        fig_num = 0
        for para in _orig_doc.paragraphs:
            for run in para.runs:
                drawings = run._element.findall(
                    './/{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
                for d in drawings:
                    blips = d.findall(
                        './/{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                    for b in blips:
                        embed = b.get(
                            '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if embed and embed in _orig_doc.part.rels:
                            fig_num += 1
                            blob = _orig_doc.part.rels[embed].target_part.blob
                            _orig_images[fig_num] = blob
                            target = _orig_doc.part.rels[embed].target_ref
                            print(f"  Extracted Figure {fig_num}: {target} ({len(blob):,} bytes)")

        _orig_tables = _orig_doc.tables
        print(f"  Found {len(_orig_tables)} tables in original docx")
        print(f"  Found {len(_orig_images)} images in original docx")
    except Exception as e:
        print(f"WARNING: Could not load original docx: {e}")
        _orig_doc = None
        _orig_images = {}
        _orig_tables = []


# ── Helper functions ─────────────────────────────────────────────────

def set_run_font(run, size=12, bold=False, italic=False, name="Times New Roman"):
    """Set font properties on a run."""
    run.font.name = name
    run.font.size = Pt(size)
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True
    # Ensure consistent font across locales
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(
            f'<w:rFonts {nsdecls("w")} w:ascii="{name}" w:hAnsi="{name}" w:cs="{name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:ascii'), name)
        rFonts.set(qn('w:hAnsi'), name)
        rFonts.set(qn('w:cs'), name)


def set_paragraph_spacing(para, space_before=None, space_after=None,
                          line_spacing=LINE_SPACING, first_indent=None,
                          left_indent=None, alignment=None):
    """Set paragraph spacing and indentation."""
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    if space_before is not None:
        pf.space_before = space_before
    else:
        pf.space_before = Pt(0)
    if space_after is not None:
        pf.space_after = space_after
    else:
        pf.space_after = Pt(0)
    if first_indent is not None:
        pf.first_line_indent = first_indent
    if left_indent is not None:
        pf.left_indent = left_indent
    if alignment is not None:
        pf.alignment = alignment


def add_empty_paragraph(doc, count=1):
    """Add empty paragraphs with proper formatting."""
    for _ in range(count):
        para = doc.add_paragraph()
        run = para.add_run("")
        set_run_font(run)
        set_paragraph_spacing(para)


def add_simple_paragraph(doc, text, bold=False, italic=False, alignment=None,
                         first_indent=None, space_before=None, space_after=None,
                         font_size=12):
    """Add a simple single-style paragraph."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, size=font_size, bold=bold, italic=italic)
    set_paragraph_spacing(para, space_before=space_before, space_after=space_after,
                          first_indent=first_indent, alignment=alignment)
    return para


def parse_inline_formatting(text):
    """Parse markdown inline formatting into segments.
    Returns list of (text, bold, italic) tuples.
    Handles ***bold-italic***, **bold**, *italic* (non-greedy).
    """
    segments = []
    # Pattern: ***bold-italic***, **bold**, *italic*
    pattern = re.compile(r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*)')
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            segments.append((text[pos:match.start()], False, False))
        if match.group(2):  # ***bold-italic***
            segments.append((match.group(2), True, True))
        elif match.group(3):  # **bold**
            segments.append((match.group(3), True, False))
        elif match.group(4):  # *italic*
            segments.append((match.group(4), False, True))
        pos = match.end()
    if pos < len(text):
        segments.append((text[pos:], False, False))
    return segments


def normalize_dashes(text):
    """Replace markdown dashes with proper unicode dashes."""
    # " -- " -> em-dash with spaces
    text = text.replace(" -- ", " \u2014 ")
    # Remaining "--" -> en-dash (common in date ranges like 2022--2024)
    text = text.replace("--", "\u2013")
    return text


def add_formatted_paragraph(doc, text, first_indent=BODY_INDENT,
                            alignment=None, space_before=None, space_after=None,
                            font_size=12, left_indent=None):
    """Add a paragraph with inline markdown formatting (bold/italic)."""
    text = normalize_dashes(text)

    para = doc.add_paragraph()
    set_paragraph_spacing(para, space_before=space_before, space_after=space_after,
                          first_indent=first_indent, alignment=alignment,
                          left_indent=left_indent)

    segments = parse_inline_formatting(text)
    for seg_text, bold, italic in segments:
        run = para.add_run(seg_text)
        set_run_font(run, size=font_size, bold=bold, italic=italic)

    return para


def add_figure(doc, fig_num):
    """Insert a figure with caption. Uses files from FIG_DIR first, then original docx."""
    if fig_num not in FIGURES:
        print(f"  WARNING: Figure {fig_num} not in FIGURES mapping")
        return

    filename, caption = FIGURES[fig_num]
    filepath = os.path.join(FIG_DIR, filename)

    # Figure number label (bold)
    para_label = doc.add_paragraph()
    run_label = para_label.add_run(f"Figure {fig_num}")
    set_run_font(run_label, bold=True)
    set_paragraph_spacing(para_label, space_before=Pt(12), space_after=Pt(0))

    # Insert image centered
    para_img = doc.add_paragraph()
    set_paragraph_spacing(para_img, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                          space_before=Pt(6), space_after=Pt(6))
    run_img = para_img.add_run()

    inserted = False
    if os.path.exists(filepath):
        try:
            run_img.add_picture(filepath, width=Inches(5.5))
            inserted = True
        except Exception as e:
            print(f"  WARNING: Could not insert figure file {fig_num}: {e}")

    if not inserted and _orig_images and fig_num in _orig_images:
        try:
            image_stream = io.BytesIO(_orig_images[fig_num])
            run_img.add_picture(image_stream, width=Inches(5.5))
            inserted = True
            print(f"  Used original docx image for Figure {fig_num}")
        except Exception as e:
            print(f"  WARNING: Could not insert original image for Figure {fig_num}: {e}")

    if not inserted:
        para_img.clear()
        run_img = para_img.add_run(f"[Figure {fig_num} not found: {filename}]")
        set_run_font(run_img, italic=True)
        print(f"  WARNING: Figure {fig_num} could not be inserted")

    # Caption (italic)
    caption_text = caption.replace(f"Figure {fig_num}. ", "")
    para_cap = doc.add_paragraph()
    run_cap = para_cap.add_run(caption_text)
    set_run_font(run_cap, size=10, italic=True)
    set_paragraph_spacing(para_cap, space_before=Pt(0), space_after=Pt(12))


def copy_table_from_original(doc, table_idx):
    """Copy a table from the original docx by index (0-based).
    Returns True if successful, False otherwise.
    """
    if _orig_tables is None or table_idx >= len(_orig_tables):
        print(f"  WARNING: Original table index {table_idx} not available")
        return False

    try:
        orig_table = _orig_tables[table_idx]
        num_rows = len(orig_table.rows)
        num_cols = len(orig_table.columns)

        new_table = doc.add_table(rows=num_rows, cols=num_cols)
        new_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        new_table.autofit = True

        # Copy table-level properties (including borders)
        orig_tbl = orig_table._tbl
        new_tbl = new_table._tbl

        orig_tblPr = orig_tbl.find(qn('w:tblPr'))
        if orig_tblPr is not None:
            new_tblPr = new_tbl.find(qn('w:tblPr'))
            if new_tblPr is not None:
                new_tbl.remove(new_tblPr)
            new_tbl.insert(0, copy.deepcopy(orig_tblPr))

        # Copy cell contents, formatting, and properties
        for r_idx, orig_row in enumerate(orig_table.rows):
            for c_idx, orig_cell in enumerate(orig_row.cells):
                new_cell = new_table.rows[r_idx].cells[c_idx]
                new_cell.text = ""

                for p_idx, orig_para in enumerate(orig_cell.paragraphs):
                    if p_idx == 0:
                        new_para = new_cell.paragraphs[0]
                    else:
                        new_para = new_cell.add_paragraph()

                    new_para._element.clear()
                    for child in orig_para._element:
                        new_para._element.append(copy.deepcopy(child))

                    orig_pPr = orig_para._element.find(qn('w:pPr'))
                    if orig_pPr is not None:
                        existing_pPr = new_para._element.find(qn('w:pPr'))
                        if existing_pPr is not None:
                            new_para._element.remove(existing_pPr)
                        new_para._element.insert(0, copy.deepcopy(orig_pPr))

                orig_tcPr = orig_cell._tc.find(qn('w:tcPr'))
                if orig_tcPr is not None:
                    existing_tcPr = new_cell._tc.find(qn('w:tcPr'))
                    if existing_tcPr is not None:
                        new_cell._tc.remove(existing_tcPr)
                    new_cell._tc.insert(0, copy.deepcopy(orig_tcPr))

            # Copy row properties
            orig_trPr = orig_row._tr.find(qn('w:trPr'))
            if orig_trPr is not None:
                new_trPr = new_table.rows[r_idx]._tr.find(qn('w:trPr'))
                if new_trPr is not None:
                    new_table.rows[r_idx]._tr.remove(new_trPr)
                new_table.rows[r_idx]._tr.insert(0, copy.deepcopy(orig_trPr))

        return True
    except Exception as e:
        print(f"  WARNING: Could not copy table {table_idx}: {e}")
        return False


def build_apa_table(doc, headers, rows, note=None):
    """Build an APA-style table from data."""
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Remove all borders, then add APA-style borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    # Header row
    header_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Top and bottom borders for header row
    for cell in header_row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            '  <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
            '  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
            '</w:tcBorders>'
        )
        tcPr.append(tcBorders)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx in range(num_cols):
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            cell_text = str(row_data[c_idx]) if c_idx < len(row_data) else ""
            # Parse inline formatting (e.g., **bold** for selected rows)
            segments = parse_inline_formatting(cell_text)
            for seg_text, seg_bold, seg_italic in segments:
                run = p.add_run(seg_text)
                set_run_font(run, size=10, bold=seg_bold, italic=seg_italic)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Bottom border on last row
    last_row = table.rows[-1]
    for cell in last_row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            '  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
            '</w:tcBorders>'
        )
        tcPr.append(tcBorders)

    # Table note
    if note:
        para_note = doc.add_paragraph()
        run_label = para_note.add_run("Note. ")
        set_run_font(run_label, size=10, italic=True)
        segments = parse_inline_formatting(note)
        for seg_text, seg_bold, seg_italic in segments:
            run = para_note.add_run(seg_text)
            set_run_font(run, size=10, italic=seg_italic)
        para_note.paragraph_format.space_before = Pt(2)
        para_note.paragraph_format.space_after = Pt(12)
        para_note.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        para_note.paragraph_format.line_spacing = LINE_SPACING

    return table


def parse_markdown_table(lines, start_idx):
    """Parse a markdown table starting at line start_idx.
    Returns (headers, rows, end_idx).
    """
    headers = []
    rows = []
    i = start_idx

    # Header row
    if i < len(lines) and lines[i].strip().startswith("|"):
        cells = [c.strip() for c in lines[i].strip().split("|")]
        cells = [c for c in cells if c]
        headers = [re.sub(r'\*\*(.+?)\*\*', r'\1', c) for c in cells]
        i += 1

    # Separator line
    if i < len(lines) and re.match(r'^[\s|:\-]+$', lines[i].strip()):
        i += 1

    # Data rows
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [c.strip() for c in lines[i].strip().split("|")]
        cells = [c for c in cells if c]
        rows.append(cells)
        i += 1

    return headers, rows, i


def add_table_caption(doc, tbl_num, caption_text):
    """Add a table caption in standard format: 'Table N.' (bold) + caption (italic)."""
    para = doc.add_paragraph()
    run_num = para.add_run(f"Table {tbl_num}.")
    set_run_font(run_num, bold=True)
    if caption_text:
        run_text = para.add_run(f" {caption_text}")
        set_run_font(run_text, italic=True)
    set_paragraph_spacing(para, space_before=Pt(12), space_after=Pt(6))
    return para


def add_table_note(doc, note_text):
    """Add a table note paragraph."""
    para = doc.add_paragraph()
    run_label = para.add_run("Note. ")
    set_run_font(run_label, size=10, italic=True)
    segments = parse_inline_formatting(note_text)
    for seg_text, seg_bold, seg_italic in segments:
        run = para.add_run(seg_text)
        set_run_font(run, size=10, italic=seg_italic)
    set_paragraph_spacing(para, space_before=Pt(2), space_after=Pt(12))
    return para


# ── Table data for Tables 1-7 (fallback if original copy fails) ─────

TABLE_CAPTIONS = {
    1: "Latent class model fit indices for Form A and Form B (W132)",
    2: "Cross-form profile distances (Euclidean distance in shared-indicator space)",
    3: "Multinomial logistic regression model comparison",
    4: "Average marginal effects on class membership probability from full multinomial logistic model",
    5: "Indirect effects of SES on class membership through AI awareness (bootstrap, 1,000 replications)",
    6: "Education coefficients (log-odds) on Anxious class membership from multi-group multinomial logistic regression, by race/ethnicity",
    7: "Cross-wave LCA validation: BIC-optimal two-class solutions across survey waves",
    8: "Sensitivity analysis: LCA model fit with and without awareness indicator",
    9: "Cross-tabulation of original four-class (with awareness) and sensitivity three-class (without awareness) assignments, Form A",
}

# Original docx table index mapping (0-based)
ORIG_TABLE_IDX = {1: 0, 2: 1, 3: 2, 4: 3, 5: 4, 6: 5, 7: 6}

# Which figures follow which tables in the original layout
TABLE_FIGURE_MAP = {
    1: [1],  # Figure 1 after Table 1
    4: [2],  # Figure 2 after Table 4
    5: [3],  # Figure 3 after Table 5
    6: [4],  # Figure 4 after Table 6
    7: [5, 6],  # Figures 5, 6 after Table 7
}


def insert_table(doc, tbl_num, md_lines=None, md_start=None):
    """Insert a table with caption. For Tables 1-7, copies from original.
    For Tables 8-9, builds from markdown.
    Returns end_idx for markdown parsing.
    """
    # Add caption
    caption = TABLE_CAPTIONS.get(tbl_num, "")
    add_table_caption(doc, tbl_num, caption)

    end_idx = md_start

    # For Tables 1-7, try to copy from original
    if tbl_num in ORIG_TABLE_IDX:
        success = copy_table_from_original(doc, ORIG_TABLE_IDX[tbl_num])
        if success:
            print(f"    Copied Table {tbl_num} from original docx")
        else:
            # Fall back to building from known data
            _build_table_fallback(doc, tbl_num)
            print(f"    Built Table {tbl_num} from fallback data")
    elif md_lines is not None and md_start is not None:
        # Build from markdown
        i = md_start
        while i < len(md_lines) and not md_lines[i].strip():
            i += 1
        if i < len(md_lines) and md_lines[i].strip().startswith("|"):
            headers, rows, end_idx = parse_markdown_table(md_lines, i)
            if headers and rows:
                build_apa_table(doc, headers, rows)
                print(f"    Built Table {tbl_num} from markdown")
        else:
            _build_table_fallback(doc, tbl_num)
            print(f"    Built Table {tbl_num} from fallback data")

    return end_idx


def _build_table_fallback(doc, tbl_num):
    """Build a table from hardcoded data when original copy fails."""
    if tbl_num == 1:
        build_apa_table(doc,
            ["Form", "Classes", "Log-Lik", "AIC", "BIC", "Entropy", "Min Class %"],
            [
                ["A", "2", "\u221219,367.5", "38,769.0", "38,881.0", "0.787", "23.8%"],
                ["A", "3", "\u221219,192.2", "38,436.3", "38,607.6", "0.656", "9.0%"],
                ["A", "4*", "\u221219,119.1", "38,308.2", "38,538.8", "0.659", "9.0%"],
                ["A", "5", "\u221219,101.5", "38,290.9", "38,580.8", "0.749", "7.9%"],
                ["B", "2", "\u221219,267.4", "38,568.7", "38,680.6", "0.723", "28.2%"],
                ["B", "3", "\u221219,107.6", "38,267.3", "38,438.3", "0.723", "14.2%"],
                ["B", "4*", "\u221219,021.8", "38,113.7", "38,343.9", "0.643", "10.6%"],
                ["B", "5", "\u221218,999.1", "38,086.1", "38,375.6", "0.596", "8.5%"],
            ],
            note="* = selected model. Selection criteria: lowest BIC with entropy > 0.6 and minimum class > 5%.")
    elif tbl_num == 2:
        build_apa_table(doc,
            ["Form A Class", "Form B Match", "Distance", "Label"],
            [
                ["A_C1", "B_C2", "0.076", "AI-Anxious"],
                ["A_C4", "B_C1", "0.071", "AI-Ambivalent"],
                ["A_C3", "B_C3", "0.313", "AI-Advantaged"],
                ["A_C2", "B_C4", "0.240", "AI-Uninformed"],
            ],
            note="Lower values indicate greater similarity between matched class profiles.")
    elif tbl_num == 3:
        build_apa_table(doc,
            ["Model", "Log-Lik", "df", "AIC", "BIC", "LRT \u03c7\u00b2", "LRT df", "LRT p"],
            [
                ["Base (demographics + race)", "\u22126,379.1", "30", "12,818.3", "13,014.5", "\u2014", "\u2014", "\u2014"],
                ["+ SES", "\u22126,313.8", "42", "12,711.6", "12,986.3", "131", "12", "< .001"],
                ["+ Awareness", "\u22124,767.9", "33", "9,601.9", "9,817.7", "3,222", "3", "< .001"],
                ["Full (all predictors)", "\u22124,746.8", "45", "9,583.6", "9,877.9", "3,265", "15", "< .001"],
            ],
            note="Reference class: Ambivalent (largest class). All models estimated on Form A (N = 5,368).")
    elif tbl_num == 4:
        build_apa_table(doc,
            ["Predictor", "Anxious", "Uninformed", "Advantaged", "Ambivalent"],
            [
                ["AI awareness (1-unit)", "\u22120.198***", "\u22120.502***", "+1.669***", "\u22120.970***"],
                ["HS or less", "+0.046***", "\u22120.011", "+0.002", "\u22120.037*"],
                ["Some college", "+0.022*", "\u22120.011", "+0.004", "\u22120.016"],
                ["Lower income", "+0.015", "+0.000", "+0.006", "\u22120.021"],
                ["Upper income", "\u22120.015", "+0.027", "+0.033**", "\u22120.045*"],
                ["Hispanic", "\u22120.023*", "+0.033", "\u22120.011", "+0.001"],
                ["Black NH", "+0.014", "\u22120.015", "\u22120.016", "+0.017"],
                ["Asian NH", "\u22120.032*", "+0.046", "+0.063***", "\u22120.077**"],
                ["Republican", "+0.062***", "\u22120.055***", "\u22120.019*", "+0.012"],
                ["Female", "+0.021**", "\u22120.008", "\u22120.037***", "+0.024"],
                ["Age 65+", "\u22120.048***", "+0.098***", "+0.024", "\u22120.074***"],
            ],
            note="Reference categories: College graduate+ (education), Middle (income), White NH (race), "
                 "30\u201349 (age), Male (gender), Democrat/Lean Dem (party). *p < .05. **p < .01. ***p < .001.")
    elif tbl_num == 5:
        build_apa_table(doc,
            ["SES Predictor", "Class Outcome", "Indirect Effect (\u03b2)", "95% CI", "p"],
            [
                ["HS or less", "AI-Advantaged", "\u22120.103", "[\u22120.123, \u22120.085]", "< .001"],
                ["HS or less", "AI-Uninformed", "+0.115", "[0.094, 0.137]", "< .001"],
                ["HS or less", "AI-Anxious", "\u22120.005", "[\u22120.009, \u22120.001]", ".013"],
                ["Upper income", "AI-Advantaged", "+0.041", "[0.027, 0.055]", "< .001"],
                ["Upper income", "AI-Uninformed", "\u22120.046", "[\u22120.062, \u22120.030]", "< .001"],
            ],
            note="Indirect effects computed as the product of coefficients with bias-corrected bootstrap confidence intervals.")
    elif tbl_num == 6:
        build_apa_table(doc,
            ["Race/Ethnicity", "Education Level", "Log-Odds (Anxious)", "SE", "p"],
            [
                ["White NH", "HS or less", "+0.386", "0.168", ".022"],
                ["Black NH", "HS or less", "+1.096", "0.499", ".028"],
                ["Hispanic", "Some college", "\u22120.234", "0.448", ".602"],
                ["Hispanic", "HS or less", "+0.307", "0.402", ".445"],
            ],
            note="Education-by-race interaction LRT: p = .007.")
    elif tbl_num == 7:
        build_apa_table(doc,
            ["Wave", "Date", "N", "BIC (2-class)", "Class Labels"],
            [
                ["W119", "Dec 2022", "10,906", "41,667.3", "AI-Optimistic + Ambivalent"],
                ["W132", "Aug 2023", "5,368", "19,259.0", "AI-Skeptic + Ambivalent"],
                ["W152", "Aug 2024", "5,363", "19,163.5", "Ambivalent + AI-Skeptic"],
            ],
            note="Cross-wave chi-square test of class proportions: \u03c7\u00b2 = 12,133.6, df = 4, p < .001.")
    elif tbl_num == 8:
        build_apa_table(doc,
            ["Form", "Indicators", "Classes", "BIC", "Entropy", "Min Class (%)"],
            [
                ["A", "6 (with awareness)", "3", "38,607.6", "0.656", "9.0"],
                ["A", "**6 (with awareness)**", "**4**", "**38,538.8**", "**0.659**", "**9.0**"],
                ["A", "**5 (no awareness)**", "**3**", "**29,253.7**", "**0.663**", "**9.0**"],
                ["A", "5 (no awareness)", "4", "29,295.3", "0.689", "2.9"],
                ["B", "6 (with awareness)", "3", "38,438.3", "0.723", "14.2"],
                ["B", "**6 (with awareness)**", "**4**", "**38,343.9**", "**0.643**", "**10.6**"],
                ["B", "**5 (no awareness)**", "**3**", "**29,338.2**", "**0.601**", "**9.8**"],
                ["B", "5 (no awareness)", "4", "29,364.2", "0.704", "2.1"],
            ],
            note="Model selection criteria: lowest BIC among models with entropy > 0.6 and minimum class size > 5%. "
                 "With six indicators (including awareness), the four-class solution is optimal for both forms. "
                 "With five indicators (awareness removed), the three-class solution is optimal for both forms. "
                 "Bold rows indicate selected models.")
    elif tbl_num == 9:
        build_apa_table(doc,
            ["Original Class (6 indicators)", "Sensitivity Class 1", "Sensitivity Class 2", "Sensitivity Class 3"],
            [
                ["Class 1: AI-Anxious (9.0%)", "**481 (99.6%)**", "2 (0.4%)", "0 (0.0%)"],
                ["Class 2: AI-Uninformed (33.1%)", "0 (0.0%)", "11 (0.6%)", "**1,764 (99.4%)**"],
                ["Class 3: AI-Advantaged (20.5%)", "0 (0.0%)", "0 (0.0%)", "**1,100 (100.0%)**"],
                ["Class 4: AI-Ambivalent (37.4%)", "0 (0.0%)", "**1,921 (95.6%)**", "89 (4.4%)"],
            ],
            note="Sensitivity Class 1 corresponds to the original AI-Anxious profile. "
                 "Sensitivity Class 2 corresponds to the original AI-Ambivalent profile. "
                 "Sensitivity Class 3 merges the original AI-Uninformed and AI-Advantaged profiles, "
                 "which share similar domain-level attitude patterns but differ primarily in awareness level. "
                 "Bold entries indicate the dominant mapping (>95% of row).")


# ── Main document builder ────────────────────────────────────────────

def build_document():
    """Build the complete Word document from the humanized markdown."""
    print("Loading original docx for figure/table extraction...")
    load_original_docx()

    print("\nReading markdown...")
    with open(MD_PATH, "r", encoding="utf-8") as f:
        content = f.read()
    lines = content.split("\n")

    print("Building document...\n")
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = MARGIN
    section.right_margin = MARGIN

    # ── Default style ──
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = LINE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    # ══════════════════════════════════════════════════════════════
    # TITLE PAGE
    # ══════════════════════════════════════════════════════════════
    print("  [1/5] Title page")

    add_empty_paragraph(doc, 3)

    title_text = ("The AI Divide: Latent Attitude Profiles and the Stratification\n"
                  "of Public Orientations Toward Artificial Intelligence")
    add_simple_paragraph(doc, title_text, bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_after=Pt(12))

    add_simple_paragraph(doc, "[Author Name]",
                         alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_simple_paragraph(doc, "[Department, University]",
                         alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_after=AUTHOR_SPACE)

    add_simple_paragraph(doc, "Author Note", bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=AUTHOR_SPACE, space_after=Pt(12))

    add_formatted_paragraph(
        doc,
        "Correspondence should be addressed to [Author information]. "
        "Data from the Pew Research Center American Trends Panel are publicly available. "
        "Replication code and materials are available from the authors upon request.",
        first_indent=BODY_INDENT)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # ABSTRACT PAGE
    # ══════════════════════════════════════════════════════════════
    print("  [2/5] Abstract page")

    add_simple_paragraph(doc, "Abstract", bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Extract abstract and keywords from markdown
    abstract_text = ""
    keywords_text = ""
    i = 0
    while i < len(lines):
        if lines[i].strip() == "**Abstract**":
            i += 1
            while i < len(lines) and not lines[i].strip():
                i += 1
            abstract_lines = []
            while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith("**Keywords"):
                abstract_lines.append(lines[i].strip())
                i += 1
            abstract_text = " ".join(abstract_lines)
            while i < len(lines):
                if lines[i].strip().startswith("**Keywords"):
                    kw_match = re.match(r'\*\*Keywords:\*\*\s*(.*)', lines[i].strip())
                    if kw_match:
                        keywords_text = kw_match.group(1)
                    break
                i += 1
            break
        i += 1

    abstract_text = normalize_dashes(abstract_text)
    para_abs = doc.add_paragraph()
    run_abs = para_abs.add_run(abstract_text)
    set_run_font(run_abs)
    set_paragraph_spacing(para_abs, first_indent=Emu(0))

    # Keywords
    para_kw = doc.add_paragraph()
    set_paragraph_spacing(para_kw, first_indent=BODY_INDENT)
    run_kw_label = para_kw.add_run("Keywords: ")
    set_run_font(run_kw_label, italic=True)
    run_kw = para_kw.add_run(keywords_text)
    set_run_font(run_kw, italic=True)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # BODY TEXT (Introduction through Conclusion)
    # ══════════════════════════════════════════════════════════════
    print("  [3/5] Body text")

    # Find section boundaries
    intro_match = re.search(r'^## 1\. Introduction', content, re.MULTILINE)
    refs_match = re.search(r'^## References', content, re.MULTILINE)
    tables_match = re.search(r'^## Tables and Figures', content, re.MULTILINE)

    if intro_match and refs_match:
        body = content[intro_match.start():refs_match.start()]
    else:
        print("  WARNING: Could not find section boundaries")
        body = content

    body_lines = body.split("\n")

    # Regex patterns
    figure_pattern = re.compile(r'\[Figure (\d+) about here.*?\]')
    table_pattern = re.compile(r'\[Table (\d+) about here.*?\]')

    inserted_tables = set()
    inserted_figures = set()

    i = 0
    while i < len(body_lines):
        line = body_lines[i].rstrip()

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # Skip horizontal rules
        if line.strip() == "---":
            i += 1
            continue

        # Section heading: ## N. Title
        sec_match = re.match(r'^## (\d+\. .+)$', line)
        if sec_match:
            add_simple_paragraph(doc, sec_match.group(1), bold=True,
                                 alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                 space_before=HEADING_SPACE)
            i += 1
            continue

        # Subsection heading: ### N.N Title
        subsec_match = re.match(r'^### (\d+\.\d+ .+)$', line)
        if subsec_match:
            add_simple_paragraph(doc, subsec_match.group(1), bold=True,
                                 alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                 space_before=HEADING_SPACE)
            i += 1
            continue

        # Figure placeholder: [Figure N about here...]
        fig_match = figure_pattern.search(line)
        if fig_match:
            fig_num = int(fig_match.group(1))
            if fig_num not in inserted_figures:
                add_figure(doc, fig_num)
                inserted_figures.add(fig_num)
            i += 1
            continue

        # Table placeholder: [Table N about here...]
        tbl_match = table_pattern.search(line)
        if tbl_match:
            tbl_num = int(tbl_match.group(1))
            if tbl_num not in inserted_tables:
                insert_table(doc, tbl_num)
                inserted_tables.add(tbl_num)
                # Insert associated figures after the table
                if tbl_num in TABLE_FIGURE_MAP:
                    for fnum in TABLE_FIGURE_MAP[tbl_num]:
                        if fnum not in inserted_figures:
                            add_figure(doc, fnum)
                            inserted_figures.add(fnum)
            i += 1
            continue

        # Inline table caption: **Table N.** (for new tables 8, 9 in body)
        tcap_match = re.match(r'^\*\*Table (\d+)\.\*\*\s*(.*)', line)
        if tcap_match:
            tbl_num = int(tcap_match.group(1))
            if tbl_num not in inserted_tables:
                # Look ahead for markdown table data
                i += 1
                # Skip empty lines
                while i < len(body_lines) and not body_lines[i].strip():
                    i += 1
                insert_table(doc, tbl_num, body_lines, i)
                inserted_tables.add(tbl_num)
                # Skip the markdown table rows
                while i < len(body_lines) and body_lines[i].strip().startswith("|"):
                    i += 1
                # Skip separator if present
                if i < len(body_lines) and re.match(r'^[\s|:\-]+$', body_lines[i].strip()):
                    i += 1
                # Skip note lines
                while i < len(body_lines) and body_lines[i].strip() and \
                      (body_lines[i].strip().startswith("*Note") or
                       body_lines[i].strip().startswith("Note")):
                    i += 1
            else:
                i += 1
            continue

        # Skip standalone markdown table rows (e.g., from already-processed tables)
        if line.strip().startswith("|"):
            i += 1
            continue

        # Skip standalone note/metadata lines
        if (line.strip().startswith("*Note") or
            line.strip().startswith("Note:") or
            re.match(r'^(Education-by-race|Cross-wave chi-square|\\?\*p <)', line.strip())):
            i += 1
            continue

        # Block quote
        if line.startswith("> "):
            quote_text = line[2:].strip()
            while i + 1 < len(body_lines) and body_lines[i + 1].strip().startswith("> "):
                i += 1
                quote_text += " " + body_lines[i].strip()[2:]
            add_formatted_paragraph(doc, quote_text,
                                    first_indent=BODY_INDENT,
                                    left_indent=Emu(457200))
            i += 1
            continue

        # Regular paragraph -- collect continuation lines
        para_text = line
        while (i + 1 < len(body_lines) and
               body_lines[i + 1].strip() and
               not body_lines[i + 1].startswith("#") and
               not body_lines[i + 1].startswith("[") and
               not body_lines[i + 1].strip().startswith("|") and
               not body_lines[i + 1].startswith("---") and
               not body_lines[i + 1].startswith("**Table ") and
               not body_lines[i + 1].startswith("**Figure ") and
               not body_lines[i + 1].startswith("*Note") and
               not body_lines[i + 1].startswith("Note:") and
               not body_lines[i + 1].startswith("> ") and
               not figure_pattern.search(body_lines[i + 1]) and
               not table_pattern.search(body_lines[i + 1]) and
               not re.match(r'^(Education-by-race|Cross-wave chi-square|\\?\*p <)',
                            body_lines[i + 1].strip())):
            i += 1
            para_text += " " + body_lines[i].strip()

        add_formatted_paragraph(doc, para_text, first_indent=BODY_INDENT)
        i += 1

    # ══════════════════════════════════════════════════════════════
    # REFERENCES
    # ══════════════════════════════════════════════════════════════
    print("  [4/5] References")

    doc.add_page_break()
    add_simple_paragraph(doc, "References", bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER)

    if refs_match and tables_match:
        refs_text = content[refs_match.end():tables_match.start()].strip()
    elif refs_match:
        refs_text = content[refs_match.end():].strip()
    else:
        refs_text = ""

    ref_entries = refs_text.split("\n\n")
    for ref in ref_entries:
        ref = ref.strip()
        if not ref or ref.startswith("---"):
            continue

        ref = normalize_dashes(ref)

        para = doc.add_paragraph()
        set_paragraph_spacing(para,
                              first_indent=Emu(-457200),
                              left_indent=Emu(457200))

        # Split on italic markers for journal/book names
        parts = re.split(r'(\*[^*]+\*)', ref)
        for part in parts:
            if part.startswith("*") and part.endswith("*"):
                run = para.add_run(part[1:-1])
                set_run_font(run, italic=True)
            else:
                run = para.add_run(part)
                set_run_font(run)

    # NOTE: Tables and figures are already embedded inline in the body text.
    # No end-of-document appendix needed.

    # ── Save ──
    doc.save(OUT_PATH)
    file_size = os.path.getsize(OUT_PATH)
    print(f"\nDocument saved to: {OUT_PATH}")
    print(f"File size: {file_size:,} bytes ({file_size / 1024 / 1024:.1f} MB)")
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"Total tables: {len(doc.tables)}")


def _add_table_note_from_markdown(doc, tbl_num, content):
    """Extract and add the note line for a table from the markdown's Tables and Figures section."""
    # Table-specific notes
    notes = {
        1: "* = selected model. Selection criteria: lowest BIC with entropy > 0.6 and minimum class > 5%.",
        2: "Lower values indicate greater similarity between matched class profiles.",
        3: "Reference class: Ambivalent (largest class). All models estimated on Form A (N = 5,368).",
        4: ("Reference categories: College graduate+ (education), Middle (income), White NH (race), "
            "30\u201349 (age), Male (gender), Democrat/Lean Dem (party). *p < .05. **p < .01. ***p < .001."),
        5: "Indirect effects computed as the product of coefficients with bias-corrected bootstrap confidence intervals.",
        6: "Education-by-race interaction LRT: p = .007.",
        7: "Cross-wave chi-square test of class proportions: \u03c7\u00b2 = 12,133.6, df = 4, p < .001.",
    }
    if tbl_num in notes:
        add_table_note(doc, notes[tbl_num])


if __name__ == "__main__":
    build_document()
