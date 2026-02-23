#!/usr/bin/env python3
"""
Generate APA 7th Edition Word document for Paper 2: The AI Divide
Target journal: New Media & Society (Sage)
"""

import os
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Paths ──────────────────────────────────────────────────────────
BASE = "/Users/hosung/AI_Polarization_Pew/manuscripts/paper2"
FIG_DIR = os.path.join(BASE, "output", "figures")
OUT_PATH = os.path.join(BASE, "manuscripts", "paper2_ai_divide.docx")

# Figure mapping
FIGURES = {
    1: ("fig_lca_profiles_form_a.png",
        "Figure 1. Class-conditional response probabilities for the four-class LCA solution (Form A). "
        "Bars represent the probability of each response category within each latent class. "
        "Top indicators (attitude, awareness) are shared across forms; bottom indicators (domains a\u2013d) are Form A-specific."),
    2: ("fig_sem01_ame_by_class.png",
        "Figure 2. Average marginal effects of key predictors on latent class membership probability (full model). "
        "Points represent AMEs with 95% confidence intervals. Reference categories: College graduate+, "
        "Middle income, White NH, Ages 30\u201349, Male, Democrat/Lean Dem."),
    3: ("fig_sem02_mediation_diagram.png",
        "Figure 3. Mediation path diagram: SES to AI awareness to latent class membership. "
        "Standardized coefficients shown for the AI-Advantaged class outcome. "
        "Dashed lines represent indirect paths; solid lines represent direct paths."),
    4: ("fig_sem03_predicted_probs.png",
        "Figure 4. Predicted probability of latent class membership by education level and race/ethnicity. "
        "Predictions from the full multinomial logistic model, holding other variables at reference/median values."),
    5: ("fig_cv01_profile_comparison.png",
        "Figure 5. Cross-wave profile comparison: class-conditional probabilities on shared indicators "
        "(attitude and awareness) across W119 (December 2022), W132 (August 2023), and W152 (August 2024)."),
    6: ("fig_cv03_education_gradient.png",
        "Figure 6. Education gradient in AI-Optimistic/Engaged class membership across survey waves. "
        "Proportion of each education group assigned to the most engaged class, with 95% confidence intervals."),
}

# ── Helper functions ───────────────────────────────────────────────

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_paragraph_format(para, font_name="Times New Roman", font_size=12,
                         bold=False, italic=False, alignment=None,
                         space_before=0, space_after=0, line_spacing=2.0,
                         first_indent=None, color=None):
    """Apply formatting to a paragraph."""
    pf = para.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    if alignment is not None:
        pf.alignment = alignment
    if first_indent is not None:
        pf.first_line_indent = Inches(first_indent)

    for run in para.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = RGBColor(*color)

def add_styled_paragraph(doc, text, font_size=12, bold=False, italic=False,
                         alignment=None, space_before=0, space_after=0,
                         first_indent=None, line_spacing=2.0):
    """Add a paragraph with specified styling."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic

    pf = para.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    if alignment is not None:
        pf.alignment = alignment
    if first_indent is not None:
        pf.first_line_indent = Inches(first_indent)
    return para

def add_mixed_paragraph(doc, segments, space_before=0, space_after=0,
                        first_indent=None, alignment=None, line_spacing=2.0):
    """Add a paragraph with mixed formatting (bold/italic segments).
    segments: list of (text, bold, italic) tuples
    """
    para = doc.add_paragraph()
    for text, bold, italic in segments:
        run = para.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.font.bold = bold
        run.font.italic = italic

    pf = para.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    if alignment is not None:
        pf.alignment = alignment
    if first_indent is not None:
        pf.first_line_indent = Inches(first_indent)
    return para

def add_figure(doc, fig_num):
    """Insert a figure with APA-style caption."""
    filename, caption = FIGURES[fig_num]
    filepath = os.path.join(FIG_DIR, filename)
    if not os.path.exists(filepath):
        add_styled_paragraph(doc, f"[Figure {fig_num} file not found: {filename}]",
                             italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        return

    # Add figure number label
    para_label = add_styled_paragraph(doc, f"Figure {fig_num}", bold=True,
                                       space_before=12, space_after=0)

    # Insert image centered
    para_img = doc.add_paragraph()
    para_img.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_img.paragraph_format.space_before = Pt(6)
    para_img.paragraph_format.space_after = Pt(6)
    run = para_img.add_run()
    run.add_picture(filepath, width=Inches(5.5))

    # Add caption (italic per APA 7th)
    # Split caption: "Figure X." is bold, rest is italic
    caption_text = caption.replace(f"Figure {fig_num}. ", "")
    para_cap = doc.add_paragraph()
    run_cap = para_cap.add_run(caption_text)
    run_cap.font.name = "Times New Roman"
    run_cap.font.size = Pt(10)
    run_cap.font.italic = True
    para_cap.paragraph_format.space_before = Pt(0)
    para_cap.paragraph_format.space_after = Pt(12)
    para_cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    para_cap.paragraph_format.line_spacing = 2.0

def add_apa_table(doc, title, headers, rows, note=None, col_widths=None):
    """Add an APA-style table with borders only on top, below header, and bottom."""
    # Table title (italic, per APA 7th)
    # Table number: bold
    title_parts = title.split(".", 1)
    para_title = doc.add_paragraph()
    run_num = para_title.add_run(title_parts[0] + ".")
    run_num.font.name = "Times New Roman"
    run_num.font.size = Pt(12)
    run_num.font.bold = True
    if len(title_parts) > 1:
        run_text = para_title.add_run(title_parts[1])
        run_text.font.name = "Times New Roman"
        run_text.font.size = Pt(12)
        run_text.font.italic = True
    para_title.paragraph_format.space_before = Pt(12)
    para_title.paragraph_format.space_after = Pt(6)
    para_title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    para_title.paragraph_format.line_spacing = 2.0

    # Create table
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Remove all borders first
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    # Header row
    header_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.font.bold = True
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Add top border to header row and bottom border
    for cell in header_row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            '  <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
            '  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
            '</w:tcBorders>'
        )
        tcPr.append(tcBorders)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Bottom border on last row
    last_row = table.rows[-1]
    for cell in last_row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            '  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
            '</w:tcBorders>'
        )
        tcPr.append(tcBorders)

    # Set column widths if specified
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Inches(width)

    # Table note
    if note:
        para_note = doc.add_paragraph()
        run_label = para_note.add_run("Note. ")
        run_label.font.name = "Times New Roman"
        run_label.font.size = Pt(10)
        run_label.font.italic = True
        run_text = para_note.add_run(note)
        run_text.font.name = "Times New Roman"
        run_text.font.size = Pt(10)
        para_note.paragraph_format.space_before = Pt(2)
        para_note.paragraph_format.space_after = Pt(12)
        para_note.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        para_note.paragraph_format.line_spacing = 2.0

    return table


def process_body_text(doc, text, first_indent=0.5):
    """Process body text, handling inline formatting like *italic* and bold."""
    # Split into paragraphs
    text = text.strip()
    if not text:
        return

    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 2.0
    if first_indent:
        pf.first_line_indent = Inches(first_indent)

    # Parse inline formatting: *italic*, **bold**, ***bold-italic***
    # Simple regex-based approach
    pattern = re.compile(r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*)')
    pos = 0
    for match in pattern.finditer(text):
        # Add text before match
        if match.start() > pos:
            run = para.add_run(text[pos:match.start()])
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

        if match.group(2):  # ***bold-italic***
            run = para.add_run(match.group(2))
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.italic = True
        elif match.group(3):  # **bold**
            run = para.add_run(match.group(3))
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run.font.bold = True
        elif match.group(4):  # *italic*
            run = para.add_run(match.group(4))
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run.font.italic = True

        pos = match.end()

    # Add remaining text
    if pos < len(text):
        run = para.add_run(text[pos:])
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    return para


# ── Main document generation ──────────────────────────────────────

def generate_document():
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # ── Set default font ──
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 2.0

    # ══════════════════════════════════════════════════════════════
    # TITLE PAGE (APA 7th)
    # ══════════════════════════════════════════════════════════════

    # Title (bold, centered, title case)
    for _ in range(3):
        add_styled_paragraph(doc, "", line_spacing=2.0)

    add_styled_paragraph(
        doc,
        "The AI Divide: Latent Attitude Profiles and the Stratification\n"
        "of Public Orientations Toward Artificial Intelligence",
        font_size=12, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=12
    )

    # Author name (centered)
    add_styled_paragraph(
        doc, "[Author Name]",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=0
    )

    # Affiliation (centered)
    add_styled_paragraph(
        doc, "[Department, University]",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=24
    )

    # Author note
    add_styled_paragraph(
        doc, "Author Note",
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=24, space_after=12
    )

    process_body_text(
        doc,
        "Correspondence should be addressed to [Author information]. "
        "Data from the Pew Research Center American Trends Panel are publicly available. "
        "Replication code and materials are available from the authors upon request.",
        first_indent=0.5
    )

    # Page break
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # ABSTRACT PAGE
    # ══════════════════════════════════════════════════════════════

    add_styled_paragraph(
        doc, "Abstract",
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=0
    )

    abstract_text = (
        "Public attitudes toward artificial intelligence (AI) are typically treated as a single "
        "spectrum from enthusiasm to anxiety, obscuring qualitative differences that may reproduce "
        "social stratification. Drawing on digital divide theory, we propose a three-level AI divide "
        "framework and use latent class analysis on Pew Research Center data (N = 10,749; American "
        "Trends Panel Wave 132) to identify four attitude profiles: AI-Anxious (9%), AI-Uninformed "
        "(33%), AI-Advantaged (21%), and AI-Ambivalent (37%). These profiles replicate across a "
        "split-sample design. Structural equation modeling reveals that socioeconomic status predicts "
        "class membership both directly and indirectly through AI awareness, and education effects "
        "vary by race/ethnicity (LRT p = .007), with stronger penalties for low education among "
        "Black respondents. Cross-wave validation (2022\u20132024) confirms structural stability. "
        "Unequal orientations toward AI constitute an emergent axis of digital inequality with "
        "implications for technology governance."
    )
    process_body_text(doc, abstract_text, first_indent=0)

    # Keywords
    para_kw = doc.add_paragraph()
    pf = para_kw.paragraph_format
    pf.first_line_indent = Inches(0.5)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 2.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    run_label = para_kw.add_run("Keywords: ")
    run_label.font.name = "Times New Roman"
    run_label.font.size = Pt(12)
    run_label.font.italic = True
    run_kw = para_kw.add_run(
        "artificial intelligence, digital divide, latent class analysis, "
        "public opinion, technology attitudes, social stratification"
    )
    run_kw.font.name = "Times New Roman"
    run_kw.font.size = Pt(12)
    run_kw.font.italic = True

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # BODY TEXT
    # ══════════════════════════════════════════════════════════════

    # Re-read the manuscript and parse sections
    md_path = os.path.join(BASE, "manuscripts", "paper2_ai_divide.md")
    with open(md_path, "r") as f:
        content = f.read()

    # Extract body sections (between first --- and References)
    # Find the start of Introduction
    intro_match = re.search(r'^## 1\. Introduction', content, re.MULTILINE)
    refs_match = re.search(r'^## References', content, re.MULTILINE)
    tables_match = re.search(r'^## Tables and Figures', content, re.MULTILINE)

    if intro_match and refs_match:
        body = content[intro_match.start():refs_match.start()]
    else:
        body = content

    # Extract references section
    if refs_match and tables_match:
        refs_text = content[refs_match.end():tables_match.start()].strip()
    elif refs_match:
        refs_text = content[refs_match.end():].strip()
    else:
        refs_text = ""

    # Process body line by line
    lines = body.split("\n")
    i = 0
    in_table = False
    table_lines = []

    # Track which figure/table placeholders we've seen
    figure_pattern = re.compile(r'\[Figure (\d+) about here.*?\]')
    table_placeholder_pattern = re.compile(r'\[Table (\d+) about here.*?\]')

    while i < len(lines):
        line = lines[i].rstrip()

        # Skip empty lines (handled by paragraph spacing)
        if not line:
            i += 1
            continue

        # Level 1 heading: ## X. Title
        if re.match(r'^## \d+\. ', line):
            heading_text = re.sub(r'^## ', '', line)
            # APA Level 1: Centered, Bold, Title Case
            add_styled_paragraph(
                doc, heading_text,
                bold=True,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_before=12, space_after=0
            )
            i += 1
            continue

        # Level 2 heading: ### X.Y Title
        if re.match(r'^### \d+\.\d+ ', line):
            heading_text = re.sub(r'^### ', '', line)
            # APA Level 2: Left-aligned, Bold
            add_styled_paragraph(
                doc, heading_text,
                bold=True,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                space_before=12, space_after=0
            )
            i += 1
            continue

        # Figure placeholder
        fig_match = figure_pattern.search(line)
        if fig_match:
            fig_num = int(fig_match.group(1))
            if fig_num in FIGURES:
                add_figure(doc, fig_num)
            i += 1
            continue

        # Table placeholder (we'll insert the table inline)
        tbl_match = table_placeholder_pattern.search(line)
        if tbl_match:
            tbl_num = int(tbl_match.group(1))
            insert_table_by_number(doc, tbl_num)
            i += 1
            continue

        # Markdown table (skip - tables are handled by placeholders or at the end)
        if line.startswith("|"):
            # Skip markdown table rows
            i += 1
            continue

        # Skip horizontal rules
        if line.startswith("---"):
            i += 1
            continue

        # Skip raw markdown bold headers like **Table X.**
        if re.match(r'^\*\*Table \d+\.', line):
            i += 1
            continue

        # Skip figure captions at the end (they're handled inline)
        if re.match(r'^\*\*Figure \d+\.', line):
            i += 1
            continue

        # Skip note lines
        if line.startswith("Note:") or line.startswith("*Author note:*") or \
           line.startswith("*Funding:*") or line.startswith("*Declaration"):
            i += 1
            continue

        # Skip standalone bold lines that are table headers
        if re.match(r'^Education-by-race interaction LRT:', line):
            i += 1
            continue

        if re.match(r'^Cross-wave chi-square test', line):
            i += 1
            continue

        # Regular paragraph
        # Collect continuation lines
        para_text = line
        while i + 1 < len(lines) and lines[i + 1].strip() and \
              not lines[i + 1].startswith("#") and \
              not lines[i + 1].startswith("[") and \
              not lines[i + 1].startswith("|") and \
              not lines[i + 1].startswith("---") and \
              not lines[i + 1].startswith("**Table") and \
              not lines[i + 1].startswith("**Figure") and \
              not figure_pattern.search(lines[i + 1]) and \
              not table_placeholder_pattern.search(lines[i + 1]):
            i += 1
            para_text += " " + lines[i].strip()

        # Clean up markdown formatting for processing
        # Remove leading ** for hypothesis labels etc.
        process_body_text(doc, para_text, first_indent=0.5)
        i += 1

    # ══════════════════════════════════════════════════════════════
    # REFERENCES
    # ══════════════════════════════════════════════════════════════

    doc.add_page_break()
    add_styled_paragraph(
        doc, "References",
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=0
    )

    # Parse references
    ref_entries = refs_text.split("\n\n")
    for ref in ref_entries:
        ref = ref.strip()
        if not ref or ref.startswith("---"):
            continue
        # Clean markdown italics
        ref_clean = ref.replace("*", "")
        para = doc.add_paragraph()
        pf = para.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 2.0
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        # Hanging indent: 0.5 inch
        pf.first_line_indent = Inches(-0.5)
        pf.left_indent = Inches(0.5)

        # Try to italicize journal names (text between periods after year)
        # Simple approach: find text in original with * markers
        parts = re.split(r'(\*[^*]+\*)', ref)
        for part in parts:
            if part.startswith("*") and part.endswith("*"):
                run = para.add_run(part[1:-1])
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
                run.font.italic = True
            else:
                run = para.add_run(part)
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)

    # Save
    doc.save(OUT_PATH)
    print(f"Document saved to: {OUT_PATH}")
    print(f"Total sections: {len(doc.sections)}")
    print(f"Total paragraphs: {len(doc.paragraphs)}")


def insert_table_by_number(doc, tbl_num):
    """Insert a formatted table by its number."""

    if tbl_num == 1:
        add_apa_table(
            doc,
            "Table 1. Latent Class Model Fit Indices for Form A and Form B (W132)",
            ["Form", "Classes", "Log-Lik", "AIC", "BIC", "Entropy", "Min Class %"],
            [
                ["A", "2", "\u221219,367.5", "38,769.0", "38,881.0", "0.787", "23.8%"],
                ["A", "3", "\u221219,192.2", "38,436.3", "38,607.6", "0.656", "9.0%"],
                ["A", "4*", "\u221219,119.1", "38,308.2", "38,538.8", "0.659", "9.0%"],
                ["A", "5", "\u221219,101.5", "38,290.9", "38,580.8", "0.749", "7.9%"],
                ["B", "2", "\u221219,267.4", "38,568.7", "38,680.6", "0.723", "28.2%"],
                ["B", "3", "\u221219,107.6", "38,267.3", "38,438.3", "0.723", "14.2%"],
                ["B", "4*", "\u221219,021.8", "38,113.7", "38,343.9", "0.643", "10.6%"],
                ["B", "5", "\u221218,999.1", "38,086.1", "38,375.6", "0.596", "8.5%"],
            ],
            note="* = selected model. Selection criteria: lowest BIC with entropy > 0.6 and minimum class > 5%."
        )

    elif tbl_num == 2:
        add_apa_table(
            doc,
            "Table 2. Cross-Form Profile Distances (Euclidean Distance in Shared-Indicator Space)",
            ["Form A Class", "Form B Match", "Distance", "Label"],
            [
                ["A_C1", "B_C2", "0.076", "AI-Anxious"],
                ["A_C4", "B_C1", "0.071", "AI-Ambivalent"],
                ["A_C3", "B_C3", "0.313", "AI-Advantaged"],
                ["A_C2", "B_C4", "0.240", "AI-Uninformed"],
            ],
            note="Lower values indicate greater similarity between matched class profiles."
        )

    elif tbl_num == 3:
        add_apa_table(
            doc,
            "Table 3. Multinomial Logistic Regression Model Comparison",
            ["Model", "Log-Lik", "df", "AIC", "BIC", "LRT \u03c7\u00b2", "LRT df", "LRT p"],
            [
                ["Base (demographics + race)", "\u22126,379.1", "30", "12,818.3", "13,014.5", "\u2014", "\u2014", "\u2014"],
                ["+ SES", "\u22126,313.8", "42", "12,711.6", "12,986.3", "131", "12", "< .001"],
                ["+ Awareness", "\u22124,767.9", "33", "9,601.9", "9,817.7", "3,222", "3", "< .001"],
                ["Full (all predictors)", "\u22124,746.8", "45", "9,583.6", "9,877.9", "3,265", "15", "< .001"],
            ],
            note="Reference class: Ambivalent (largest class). All models estimated on Form A (N = 5,368)."
        )

    elif tbl_num == 4:
        add_apa_table(
            doc,
            "Table 4. Average Marginal Effects on Class Membership Probability From Full Multinomial Logistic Model",
            ["Predictor", "Anxious", "Uninformed", "Advantaged", "Ambivalent"],
            [
                ["AI awareness (1-unit)", "\u22120.198***", "\u22120.502***", "+1.669***", "\u22120.970***"],
                ["HS or less", "+0.046***", "\u22120.011", "+0.002", "\u22120.037*"],
                ["Some college", "+0.022*", "\u22120.011", "+0.004", "\u22120.016"],
                ["Lower income", "+0.015", "+0.000", "+0.006", "\u22120.021"],
                ["Upper income", "\u22120.015", "+0.027", "+0.033**", "\u22120.045*"],
                ["Hispanic", "\u22120.023*", "+0.033", "\u22120.011", "+0.001"],
                ["Black NH", "+0.014", "\u22120.015", "\u22120.016", "+0.017"],
                ["Asian NH", "\u22120.032*", "+0.046", "+0.063***", "\u22120.077**"],
                ["Republican", "+0.062***", "\u22120.055***", "\u22120.019*", "+0.012"],
                ["Female", "+0.021**", "\u22120.008", "\u22120.037***", "+0.024"],
                ["Age 65+", "\u22120.048***", "+0.098***", "+0.024", "\u22120.074***"],
            ],
            note="Reference categories: College graduate+ (education), Middle (income), White NH (race), "
                 "30\u201349 (age), Male (gender), Democrat/Lean Dem (party). "
                 "*p < .05. **p < .01. ***p < .001."
        )
        # Insert Figure 2 after Table 4
        add_figure(doc, 2)

    elif tbl_num == 5:
        add_apa_table(
            doc,
            "Table 5. Indirect Effects of SES on Class Membership Through AI Awareness (Bootstrap, 1,000 Replications)",
            ["SES Predictor", "Class Outcome", "Indirect Effect (\u03b2)", "95% CI", "p"],
            [
                ["HS or less", "AI-Advantaged", "\u22120.103", "[\u22120.123, \u22120.085]", "< .001"],
                ["HS or less", "AI-Uninformed", "+0.115", "[0.094, 0.137]", "< .001"],
                ["HS or less", "AI-Anxious", "\u22120.005", "[\u22120.009, \u22120.001]", ".013"],
                ["Upper income", "AI-Advantaged", "+0.041", "[0.027, 0.055]", "< .001"],
                ["Upper income", "AI-Uninformed", "\u22120.046", "[\u22120.062, \u22120.030]", "< .001"],
            ],
            note="Indirect effects computed as the product of coefficients with bias-corrected bootstrap confidence intervals."
        )
        # Insert Figure 3 after Table 5
        add_figure(doc, 3)

    elif tbl_num == 6:
        add_apa_table(
            doc,
            "Table 6. Education Coefficients (Log-Odds) on Anxious Class Membership From Multi-Group Multinomial Logistic Regression, by Race/Ethnicity",
            ["Race/Ethnicity", "Education Level", "Log-Odds (Anxious)", "SE", "p"],
            [
                ["White NH", "HS or less", "+0.386", "0.168", ".022"],
                ["Black NH", "HS or less", "+1.096", "0.499", ".028"],
                ["Hispanic", "Some college", "\u22120.234", "0.448", ".602"],
                ["Hispanic", "HS or less", "+0.307", "0.402", ".445"],
            ],
            note="Education-by-race interaction LRT: p = .007."
        )
        # Insert Figure 4 after Table 6
        add_figure(doc, 4)

    elif tbl_num == 7:
        add_apa_table(
            doc,
            "Table 7. Cross-Wave LCA Validation: BIC-Optimal Two-Class Solutions Across Survey Waves",
            ["Wave", "Date", "N", "BIC (2-class)", "Class Labels"],
            [
                ["W119", "Dec 2022", "10,906", "41,667.3", "AI-Optimistic + Ambivalent"],
                ["W132", "Aug 2023", "5,368", "19,259.0", "AI-Skeptic + Ambivalent"],
                ["W152", "Aug 2024", "5,363", "19,163.5", "Ambivalent + AI-Skeptic"],
            ],
            note="Cross-wave chi-square test of class proportions: \u03c7\u00b2 = 12,133.6, df = 4, p < .001."
        )
        # Insert Figures 5 and 6 after Table 7
        add_figure(doc, 5)
        add_figure(doc, 6)


if __name__ == "__main__":
    generate_document()
