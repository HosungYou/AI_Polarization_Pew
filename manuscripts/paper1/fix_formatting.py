"""
Comprehensive Word formatting fix for TF&SC (Elsevier) submission.
Fixes theme fonts, style definitions, run-level fonts, and figure sizing.
"""
from docx import Document
from docx.shared import Pt, Inches, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree
import os

DOCX_PATH = '/Users/hosung/AI_Polarization_Pew/manuscripts/paper1/paper1_revised.docx'
FONT_NAME = 'Times New Roman'
FONT_SIZE = Pt(12)
TABLE_FONT_SIZE = Pt(10)

doc = Document(DOCX_PATH)

# ============================================================
# 1. FIX THEME FONTS (root cause of Aptos appearing)
# ============================================================
print("1. Fixing theme fonts...")
import zipfile, shutil, tempfile
# Directly edit theme XML inside the docx zip
tmp_dir = tempfile.mkdtemp()
tmp_docx = os.path.join(tmp_dir, 'temp.docx')
shutil.copy2(DOCX_PATH, tmp_docx)
with zipfile.ZipFile(tmp_docx, 'r') as zin:
    namelist = zin.namelist()
    theme_file = [n for n in namelist if 'theme' in n.lower() and n.endswith('.xml')]
    if theme_file:
        theme_content = zin.read(theme_file[0]).decode('utf-8')
        # Replace all Aptos variants with Times New Roman
        for aptos in ['Aptos Display', 'Aptos', 'Calibri', 'Calibri Light']:
            theme_content = theme_content.replace(f'typeface="{aptos}"', f'typeface="{FONT_NAME}"')
        # Write modified zip
        with zipfile.ZipFile(DOCX_PATH, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in namelist:
                if item == theme_file[0]:
                    zout.writestr(item, theme_content.encode('utf-8'))
                else:
                    zout.writestr(item, zin.read(item))
        print(f"   Theme fonts replaced with {FONT_NAME}")
    else:
        print("   No theme file found")
shutil.rmtree(tmp_dir)
# Reload document after theme fix
doc = Document(DOCX_PATH)

# ============================================================
# 2. FIX DEFAULT DOCUMENT FONT
# ============================================================
print("2. Fixing document default font...")
styles_element = doc.styles.element
rPrDefault = styles_element.find(qn('w:docDefaults') + '/' + qn('w:rPrDefault') + '/' + qn('w:rPr'))
if rPrDefault is not None:
    # Remove theme font references
    for rFonts in rPrDefault.findall(qn('w:rFonts')):
        rFonts.set(qn('w:ascii'), FONT_NAME)
        rFonts.set(qn('w:hAnsi'), FONT_NAME)
        rFonts.set(qn('w:cs'), FONT_NAME)
        rFonts.set(qn('w:eastAsia'), FONT_NAME)
        # Remove theme references
        for attr in ['w:asciiTheme', 'w:hAnsiTheme', 'w:cstheme', 'w:eastAsiaTheme']:
            if rFonts.get(qn(attr)) is not None:
                del rFonts.attrib[qn(attr)]
    for sz in rPrDefault.findall(qn('w:sz')):
        sz.set(qn('w:val'), '24')  # 24 half-points = 12pt
    for sz in rPrDefault.findall(qn('w:szCs')):
        sz.set(qn('w:val'), '24')
    print(f"   Default: {FONT_NAME} 12pt")

# ============================================================
# 3. FIX ALL STYLE DEFINITIONS
# ============================================================
print("3. Fixing style definitions...")
for style in doc.styles:
    try:
        if style.font:
            style.font.name = FONT_NAME
            if style.name in ('Title', 'Heading 1', 'Heading 2', 'Heading 3',
                              'Normal', 'Body Text', 'First Paragraph',
                              'Author', 'Date', 'Subtitle', 'Abstract',
                              'Block Text', 'Caption', 'Image Caption',
                              'Table Caption', 'Figure'):
                style.font.size = FONT_SIZE

            # Remove theme font references from style XML
            rPr = style.element.find(qn('w:rPr'))
            if rPr is not None:
                for rFonts in rPr.findall(qn('w:rFonts')):
                    rFonts.set(qn('w:ascii'), FONT_NAME)
                    rFonts.set(qn('w:hAnsi'), FONT_NAME)
                    for attr in ['w:asciiTheme', 'w:hAnsiTheme', 'w:cstheme', 'w:eastAsiaTheme']:
                        if rFonts.get(qn(attr)) is not None:
                            del rFonts.attrib[qn(attr)]

        # Set line spacing for paragraph styles
        if hasattr(style, 'paragraph_format') and style.paragraph_format:
            pf = style.paragraph_format
            if style.name not in ('Caption', 'Table Caption', 'Image Caption', 'Figure',
                                   'Table Grid', 'Compact'):
                pf.line_spacing = 2.0
                pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
    except Exception as e:
        pass

print(f"   All styles set to {FONT_NAME}")

# ============================================================
# 4. FIX ALL PARAGRAPHS AND RUNS
# ============================================================
print("4. Fixing all paragraphs and runs...")
para_count = 0
for para in doc.paragraphs:
    # Set paragraph formatting
    pf = para.paragraph_format
    pf.line_spacing = 2.0
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    style_name = para.style.name if para.style else ''

    # Heading alignment
    if 'Heading 1' in style_name or 'Title' in style_name:
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.first_line_indent = Inches(0)
    elif 'Heading' in style_name:
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Inches(0)
    elif 'Author' in style_name or 'Date' in style_name:
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.first_line_indent = Inches(0)
    elif 'Caption' in style_name or 'Figure' in style_name or 'Image' in style_name:
        pf.first_line_indent = Inches(0)
    else:
        # Body paragraphs: check if it's a figure caption (starts with *Figure or Figure)
        text = para.text.strip()
        if text.startswith('Figure') or text.startswith('*Figure') or text.startswith('Table'):
            pf.first_line_indent = Inches(0)
        elif text.startswith('|') or text.startswith('Note.') or text.startswith('*Note'):
            pf.first_line_indent = Inches(0)
        elif text == '' or text.startswith('---') or text.startswith('\\newpage'):
            pf.first_line_indent = Inches(0)
        else:
            pf.first_line_indent = Inches(0.5)

    # Fix all runs
    for run in para.runs:
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE

        # Remove theme font references from run XML
        rPr = run._element.find(qn('w:rPr'))
        if rPr is not None:
            for rFonts in rPr.findall(qn('w:rFonts')):
                rFonts.set(qn('w:ascii'), FONT_NAME)
                rFonts.set(qn('w:hAnsi'), FONT_NAME)
                rFonts.set(qn('w:cs'), FONT_NAME)
                for attr in ['w:asciiTheme', 'w:hAnsiTheme', 'w:cstheme', 'w:eastAsiaTheme']:
                    if rFonts.get(qn(attr)) is not None:
                        del rFonts.attrib[qn(attr)]

        # Heading formatting
        if 'Heading 1' in style_name or 'Title' in style_name:
            run.font.bold = True
        elif 'Heading 2' in style_name:
            run.font.bold = True
        elif 'Heading 3' in style_name:
            run.font.bold = True
            run.font.italic = True

    para_count += 1

print(f"   Processed {para_count} paragraphs")

# ============================================================
# 5. FIX TABLE FORMATTING
# ============================================================
print("5. Fixing tables...")
table_count = 0
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.line_spacing = 1.0
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                para.paragraph_format.space_before = Pt(1)
                para.paragraph_format.space_after = Pt(1)
                para.paragraph_format.first_line_indent = Inches(0)
                for run in para.runs:
                    run.font.name = FONT_NAME
                    run.font.size = TABLE_FONT_SIZE
                    rPr = run._element.find(qn('w:rPr'))
                    if rPr is not None:
                        for rFonts in rPr.findall(qn('w:rFonts')):
                            rFonts.set(qn('w:ascii'), FONT_NAME)
                            rFonts.set(qn('w:hAnsi'), FONT_NAME)
                            for attr in ['w:asciiTheme', 'w:hAnsiTheme', 'w:cstheme', 'w:eastAsiaTheme']:
                                if rFonts.get(qn(attr)) is not None:
                                    del rFonts.attrib[qn(attr)]
    table_count += 1
print(f"   Processed {table_count} tables")

# ============================================================
# 6. FIX FIGURE SIZING
# ============================================================
print("6. Fixing figure sizes...")
img_count = 0
for para in doc.paragraphs:
    for run in para.runs:
        inline_shapes = run._element.findall('.//' + qn('wp:inline'))
        for inline in inline_shapes:
            extent = inline.find(qn('wp:extent'))
            if extent is not None:
                old_cx = int(extent.get('cx'))
                old_cy = int(extent.get('cy'))
                new_cx = Inches(6.0)
                ratio = new_cx / old_cx
                new_cy = int(old_cy * ratio)
                extent.set('cx', str(int(new_cx)))
                extent.set('cy', str(new_cy))

                # Also fix the graphic extent
                for graphic_extent in inline.findall('.//' + qn('a:ext')):
                    graphic_extent.set('cx', str(int(new_cx)))
                    graphic_extent.set('cy', str(new_cy))

                img_count += 1
                print(f"   Image {img_count}: {old_cx/914400:.1f}in -> 6.0in")

# ============================================================
# 7. PAGE SETUP + LINE NUMBERING
# ============================================================
print("7. Fixing page setup and line numbering...")
for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Remove any line numbering
    sectPr = section._sectPr
    for old in sectPr.findall(qn('w:lnNumType')):
        sectPr.remove(old)

print("   Letter size, 1in margins, no line numbering")

# ============================================================
# 8. SAVE
# ============================================================
doc.save(DOCX_PATH)
file_size = os.path.getsize(DOCX_PATH)
print(f"\nSaved: {DOCX_PATH} ({file_size/1024:.0f} KB)")
print("DONE - All formatting applied.")
